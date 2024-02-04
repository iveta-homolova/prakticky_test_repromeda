import docx 
  
 
doc = docx.Document() 
  

doc.add_heading('GeeksForGeeks', 0) 
doc.add_run('bold').bold = True
doc.alignment = 1
  
 
data = ( 
    ("Jméno a příjmení:", 'argument 1'), 
    ("Datum odběru:", 'argument 2'), 
    ("Rodné čislo:", 'argument 3') 
) 
  

table = doc.add_table(rows=3, cols=2) 

row = table.rows[0].cells 
row[0].text = "Jméno a příjmení:"
row[1].text = "Datum odběru:"
row[1].text = "Rodné čislo:"


for  Jméno a příjmení,  Datum odběru, Rodné čislo in data: 
    row = table.add_row().cells 
    row[0].text =  input("argument1:")
    row[1].text = input("argument2:") 
    row[2].text = input("argument3:") 
  

doc.save('demo.docx')

